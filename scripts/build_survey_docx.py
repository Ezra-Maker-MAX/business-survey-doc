"""
通用化业务调研问卷 Word 文档生成器
- 基于 V4.0 EAM 调研问卷模板（已验证有效的格式）
- 支持参数化配置：调研对象、业务范围、问题列表
- 严格遵循 V4.0 文档结构（9章正文+5附录+签字栏）

使用示例：
    from build_survey_docx import generate_survey_docx

    config = {
        'project_name': 'XXX项目',
        'doc_number': 'XXX_V1.0_20260812',
        'version': 'V1.0',
        ...
    }
    generate_survey_docx(config, 'output.docx')
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# 基础辅助函数
# ============================================================

def set_cell_bg(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    cell.text = ''
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1, center=True):
    sizes = {0: 22, 1: 16, 2: 14, 3: 12}
    p = doc.add_paragraph()
    if center and level <= 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8 if level <= 1 else 6)
    p.paragraph_format.space_after = Pt(6 if level <= 1 else 4)
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True
    if level == 0:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def add_para(doc, text='', bold=False, size=11, color=None, center=False, indent_first=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    if text:
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = color
    return p


def add_mixed(doc, parts, indent_first=False):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3
    for text, fmt in parts:
        run = p.add_run(text)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(fmt.get('size', 11))
        if fmt.get('bold'):
            run.bold = True
        if fmt.get('color'):
            run.font.color.rgb = fmt['color']
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run('')
    run.text = text
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run.font.size = Pt(size)


def set_table_widths(table, widths_cm):
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_cm):
                cell.width = Cm(widths_cm[idx])


def style_header_row(table, bg='1F4E79', font_color=None):
    if font_color is None:
        font_color = RGBColor(0xFF, 0xFF, 0xFF)
    for cell in table.rows[0].cells:
        set_cell_bg(cell, bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = font_color
                run.bold = True


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement('w:' + border_name)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '7F7F7F')
        tblBorders.append(b)
    tblPr.append(tblBorders)


def render_question(doc, q_item):
    level_color_map = {
        '必问': RGBColor(0xC0, 0x00, 0x00),
        '应问': RGBColor(0xED, 0x7D, 0x31),
        '深挖': RGBColor(0x70, 0xAD, 0x47),
    }
    color = level_color_map.get(q_item['level'], RGBColor(0x1F, 0x4E, 0x79))

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7)
    p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.3

    run_idx = p.add_run('Q' + str(q_item['idx']) + '  ')
    run_idx.font.name = 'Microsoft YaHei'
    run_idx._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_idx.bold = True
    run_idx.font.size = Pt(11)
    run_idx.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    run_level = p.add_run('[' + q_item['level'] + ']  ')
    run_level.font.name = 'Microsoft YaHei'
    run_level._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_level.bold = True
    run_level.font.size = Pt(10)
    run_level.font.color.rgb = color

    run_q = p.add_run(q_item['q'])
    run_q.font.name = 'Microsoft YaHei'
    run_q._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run_q.font.size = Pt(11)

    if q_item.get('follow'):
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(0.7)
        p2.paragraph_format.space_after = Pt(2)
        rk = p2.add_run('建议追问:')
        rk.font.name = 'Microsoft YaHei'
        rk._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rk.italic = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        rv = p2.add_run(q_item['follow'])
        rv.font.name = 'Microsoft YaHei'
        rv._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rv.italic = True
        rv.font.size = Pt(10)
        rv.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    if q_item.get('need'):
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.7)
        p3.paragraph_format.space_after = Pt(2)
        rk = p3.add_run('需要查看/索取:')
        rk.font.name = 'Microsoft YaHei'
        rk._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rk.italic = True
        rk.font.size = Pt(10)
        rk.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        rv = p3.add_run(q_item['need'])
        rv.font.name = 'Microsoft YaHei'
        rv._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        rv.italic = True
        rv.font.size = Pt(10)
        rv.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    p4 = doc.add_paragraph()
    p4.paragraph_format.left_indent = Cm(0.7)
    p4.paragraph_format.space_after = Pt(8)
    r1 = p4.add_run('答复:')
    r1.font.name = 'Microsoft YaHei'
    r1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p4.add_run(' ' * 100).font.size = Pt(11)


# ============================================================
# 章节渲染函数
# ============================================================

def render_cover(doc, config):
    add_heading(doc, config['project_name'], level=0)
    add_heading(doc, config['doc_title'], level=0)
    if config.get('subtitle'):
        add_para(doc, config['subtitle'], center=True, bold=True, size=12,
                 color=RGBColor(0xC0, 0x00, 0x00))
    add_para(doc)

    info_rows = config['info_rows']
    info_table = doc.add_table(rows=len(info_rows), cols=4)
    set_table_widths(info_table, [3.0, 6.0, 3.0, 5.0])
    for r_idx, (k1, v1, k2, v2) in enumerate(info_rows):
        row = info_table.rows[r_idx]
        set_cell_text(row.cells[0], k1, bold=True, size=11)
        set_cell_bg(row.cells[0], 'F2F2F2')
        if k2 == '' and v2 == '':
            row.cells[1].merge(row.cells[3])
            set_cell_text(row.cells[1], v1, size=11)
        else:
            set_cell_text(row.cells[1], v1, size=11)
            set_cell_text(row.cells[2], k2, bold=True, size=11)
            set_cell_bg(row.cells[2], 'F2F2F2')
            set_cell_text(row.cells[3], v2, size=11)
    set_table_borders(info_table)
    doc.add_paragraph()


def render_attendees(doc, attending_dept):
    add_heading(doc, '一、参会人员', level=1)
    add_mixed(doc, [
        ('建议涵盖:', {'bold': True}),
        (f'{attending_dept};', {}),
        ('调研对象现场负责人与调研实施方共同参会。',
         {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)})
    ])

    attendee_table = doc.add_table(rows=12, cols=6)
    attendee_headers = ['序号', '部门', '姓名', '职务/角色', '联系方式', '是否参会']
    for i, h in enumerate(attendee_headers):
        set_cell_text(attendee_table.rows[0].cells[i], h, bold=True, size=11)
    for r in range(1, 12):
        set_cell_text(attendee_table.rows[r].cells[0], str(r), size=11)
    set_table_widths(attendee_table, [1.2, 3.0, 2.5, 2.5, 3.0, 2.8])
    style_header_row(attendee_table)
    set_table_borders(attendee_table)
    doc.add_paragraph()


def render_instruction(doc, project_context):
    add_heading(doc, '二、调研说明', level=1)

    add_mixed(doc, [('1. 项目背景', {'bold': True, 'color': RGBColor(0x1F, 0x4E, 0x79)})])
    add_para(doc, project_context, size=11, indent_first=True)

    add_mixed(doc, [('2. 调研目的', {'bold': True, 'color': RGBColor(0x1F, 0x4E, 0x79)})])
    add_mixed(doc, [('本次调研期望达到以下目的:', {'bold': True})], indent_first=True)
    for bullet in [
        '全面、系统、重点突出地了解调研对象在相关业务的现状、问题与需求;',
        '识别各业务环节中数字化缺失的环节、重复录入、数据孤岛与流程断点;',
        '明确系统实施范围、优先级与一/二期建设策略;',
        '加深调研团队与现场业务团队的相互了解,推动后续蓝图设计与UAT落地。',
    ]:
        add_bullet(doc, bullet)

    add_mixed(doc, [('3. 调研原则', {'bold': True, 'color': RGBColor(0x1F, 0x4E, 0x79)})])
    for bullet in [
        '先问现状,再看现场;先还原流程,再讨论系统功能;',
        '尽量拿真实表单、图纸和系统截图,不只听口述;',
        '调研结论用于指导系统一期与二期建设规划,确保需求可落地。',
    ]:
        add_bullet(doc, bullet)

    add_mixed(doc, [('4. 提问分级', {'bold': True, 'color': RGBColor(0x1F, 0x4E, 0x79)})])
    level_table = doc.add_table(rows=4, cols=2)
    level_rows = [
        ('等级', '说明'),
        ('必问', '现场必须完成的最小调研项,决定能否出具调研总结'),
        ('应问', '时间允许尽量完成,对系统设计输入与蓝图有重要补充'),
        ('深挖', '发现问题后深入追问,需要有业务背景或现场观察的引导'),
    ]
    for ri, (k, v) in enumerate(level_rows):
        set_cell_text(level_table.rows[ri].cells[0], k, size=11, bold=(ri == 0))
        set_cell_text(level_table.rows[ri].cells[1], v, size=11, bold=(ri == 0))
    set_table_widths(level_table, [2.5, 14])
    style_header_row(level_table)
    set_table_borders(level_table)
    doc.add_paragraph()


def render_toc(doc, toc_items):
    add_heading(doc, '三、调研目录', level=1)

    toc_table = doc.add_table(rows=len(toc_items) + 1, cols=4)
    toc_headers = ['模块', '业务领域', '问题编号', '问题数']
    for i, h in enumerate(toc_headers):
        set_cell_text(toc_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, (mod, area, qrange, cnt) in enumerate(toc_items):
        is_main = not mod.startswith('\u3000')
        set_cell_text(toc_table.rows[ri + 1].cells[0], mod, size=10, bold=is_main)
        set_cell_text(toc_table.rows[ri + 1].cells[1], area, size=10, bold=is_main)
        set_cell_text(toc_table.rows[ri + 1].cells[2], qrange, size=10)
        set_cell_text(toc_table.rows[ri + 1].cells[3], cnt, size=10)
    set_table_widths(toc_table, [2.0, 6.0, 5.0, 2.0])
    style_header_row(toc_table)
    set_table_borders(toc_table)
    doc.add_paragraph()


def render_overview(doc, overview_items, core_upgrade=None):
    add_heading(doc, '四、总体概述', level=1)

    for item in overview_items:
        if isinstance(item, str):
            add_para(doc, item, size=11, indent_first=True)
        elif isinstance(item, tuple) and len(item) == 2:
            parts, indent = item
            add_mixed(doc, parts, indent_first=indent)
        else:
            add_mixed(doc, item)

    if core_upgrade:
        add_mixed(doc, [
            ('\u2605 核心升级:', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
            (core_upgrade, {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
        ])

    doc.add_paragraph()


def render_current_status(doc, modules):
    add_heading(doc, '五、现状描述', level=1)

    for mod in modules:
        add_heading(doc, f"{mod['sec_num']} {mod['name']}", level=2)
        if mod.get('intro'):
            add_para(doc, mod['intro'], size=11, indent_first=True)

        if mod.get('module_note_parts'):
            add_mixed(doc, mod['module_note_parts'])

        doc.add_paragraph()

        for sub_idx, (sub_name, sub_qs) in enumerate(mod['subs'], 1):
            add_heading(doc, f"{mod['sec_num']}.{sub_idx} {sub_name}", level=3)
            if mod.get('sub_intro_for_index') and sub_idx == mod.get('sub_intro_for_index'):
                add_para(doc, mod['sub_intro_for_text'], size=10,
                         indent_first=True, color=RGBColor(0x59, 0x59, 0x59))
            for q in sub_qs:
                render_question(doc, q)
            doc.add_paragraph()


def render_high_level_interview(doc, interview_topics):
    add_heading(doc, '六、高层访谈', level=1)
    add_mixed(doc, [('面向:', {'bold': True}),
                    ('调研对象的厂领导/总监/经理', {})])
    add_para(doc, '建议围绕以下主题展开高层访谈:', size=11, indent_first=True)

    exec_table = doc.add_table(rows=len(interview_topics) + 1, cols=3)
    for i, h in enumerate(['序号', '主题', '访谈要点']):
        set_cell_text(exec_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, (num, theme, detail) in enumerate(interview_topics):
        set_cell_text(exec_table.rows[ri + 1].cells[0], num, size=11)
        set_cell_text(exec_table.rows[ri + 1].cells[1], theme, bold=True, size=11)
        set_cell_text(exec_table.rows[ri + 1].cells[2], detail, size=11)
    set_table_widths(exec_table, [1.2, 4.0, 11.0])
    style_header_row(exec_table)
    set_table_borders(exec_table)
    doc.add_paragraph()


def render_it_systems(doc, systems):
    add_heading(doc, '七、IT系统', level=1)
    add_para(doc, '了解当前相关IT系统使用情况：', size=11, indent_first=True)

    it_table = doc.add_table(rows=len(systems) + 1, cols=5)
    headers = ['系统名称', '覆盖范围', '使用部门', '是否满足需求', '备注']
    for i, h in enumerate(headers):
        set_cell_text(it_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, row_data in enumerate(systems):
        set_cell_text(it_table.rows[ri + 1].cells[0], row_data[0], size=11)
        for ci in range(1, 5):
            set_cell_text(it_table.rows[ri + 1].cells[ci], '', size=11)
    set_table_widths(it_table, [3.5, 3.5, 3.0, 3.0, 3.0])
    style_header_row(it_table)
    set_table_borders(it_table)
    doc.add_paragraph()


def render_data_interfaces(doc, flows):
    add_heading(doc, '八、数据接口', level=1)
    add_para(doc, '了解相关系统之间的数据集成需求：', size=11, indent_first=True)

    int_table = doc.add_table(rows=len(flows) + 1, cols=4)
    headers = ['数据流向', '接口方式', '同步频率', '备注']
    for i, h in enumerate(headers):
        set_cell_text(int_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, row_data in enumerate(flows):
        set_cell_text(int_table.rows[ri + 1].cells[0], row_data[0], size=11)
        for ci in range(1, 4):
            set_cell_text(int_table.rows[ri + 1].cells[ci], '', size=11)
    set_table_widths(int_table, [4.0, 4.0, 3.0, 4.0])
    style_header_row(int_table)
    set_table_borders(int_table)
    doc.add_paragraph()


def render_reports(doc, reports):
    add_heading(doc, '九、标准报表', level=1)
    add_para(doc, '了解相关报表需求：', size=11, indent_first=True)

    rpt_table = doc.add_table(rows=len(reports) + 1, cols=4)
    headers = ['报表名称', '报表内容', '生成频率', '使用部门']
    for i, h in enumerate(headers):
        set_cell_text(rpt_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, row_data in enumerate(reports):
        set_cell_text(rpt_table.rows[ri + 1].cells[0], row_data[0], size=11)
        for ci in range(1, 4):
            set_cell_text(rpt_table.rows[ri + 1].cells[ci], '', size=11)
    set_table_widths(rpt_table, [4.0, 5.0, 3.0, 3.0])
    style_header_row(rpt_table)
    set_table_borders(rpt_table)
    doc.add_paragraph()


def render_appendix_resources(doc, resources):
    add_heading(doc, '附录', level=1)
    add_heading(doc, '附录A：资料索取清单', level=2)

    res_table = doc.add_table(rows=len(resources) + 1, cols=4)
    headers = ['序号', '类别', '资料名称', '优先级']
    for i, h in enumerate(headers):
        set_cell_text(res_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, (idx, cat, name, prio) in enumerate(resources):
        set_cell_text(res_table.rows[ri + 1].cells[0], str(idx), size=10)
        set_cell_text(res_table.rows[ri + 1].cells[1], cat, size=10)
        set_cell_text(res_table.rows[ri + 1].cells[2], name, size=10)
        set_cell_text(res_table.rows[ri + 1].cells[3], prio, size=10)
        if prio == '\u9ad8':
            set_cell_bg(res_table.rows[ri + 1].cells[3], 'FCE4EC')
        elif prio == '\u4e2d':
            set_cell_bg(res_table.rows[ri + 1].cells[3], 'FFF3E0')
    set_table_widths(res_table, [1.0, 2.5, 9.0, 1.5])
    style_header_row(res_table)
    set_table_borders(res_table)
    doc.add_paragraph()


def render_appendix_sampling(doc, sampling_desc, sampling_cols=None):
    add_heading(doc, '附录B：典型对象现场抽查', level=2)
    add_para(doc, sampling_desc, size=11, indent_first=True)

    if sampling_cols is None:
        sampling_cols = ['序号', '资产编号', '资产名称', '所属模块', '业务子项',
                         '产线/区域', '类型', '投运年份', '当前状态',
                         '维护方式', '当前系统', '主要痛点', '数字化建议']

    sampling_table = doc.add_table(rows=8, cols=len(sampling_cols))
    for i, h in enumerate(sampling_cols):
        set_cell_text(sampling_table.rows[0].cells[i], h, bold=True, size=9)
    for ri in range(1, 8):
        set_cell_text(sampling_table.rows[ri].cells[0], str(ri), size=9)
        for ci in range(1, len(sampling_cols)):
            set_cell_text(sampling_table.rows[ri].cells[ci], '', size=9)
    widths = [0.8, 1.3, 1.5, 1.2, 1.5, 1.3, 1.0, 1.0, 1.2, 1.2, 1.2, 1.5, 1.5]
    set_table_widths(sampling_table, widths[:len(sampling_cols)])
    style_header_row(sampling_table)
    set_table_borders(sampling_table)
    doc.add_paragraph()


def render_appendix_pain_points(doc):
    add_heading(doc, '附录C：痛点需求池', level=2)
    PAIN_COLS = ['序号', '来源部门', '现状问题', '发生场景', '影响',
                 '根因', '用户需求', '数字化建议', '优先级', '责任系统', '备注']
    pain_table = doc.add_table(rows=11, cols=len(PAIN_COLS))
    for i, h in enumerate(PAIN_COLS):
        set_cell_text(pain_table.rows[0].cells[i], h, bold=True, size=9)
    for ri in range(1, 11):
        set_cell_text(pain_table.rows[ri].cells[0], str(ri), size=9)
        for ci in range(1, len(PAIN_COLS)):
            set_cell_text(pain_table.rows[ri].cells[ci], '', size=9)
    set_table_widths(pain_table, [0.8, 1.5, 2.0, 1.5, 1.5, 1.5, 2.0, 2.0, 1.0, 1.2, 1.0])
    style_header_row(pain_table)
    set_table_borders(pain_table)
    doc.add_paragraph()


def render_appendix_summary_dims(doc, summary_dims):
    add_heading(doc, '附录D：调研总结', level=2)
    add_para(doc, '调研完成后,请按以下维度撰写调研总结：', size=11, indent_first=True)
    for dim in summary_dims:
        add_bullet(doc, dim)
    doc.add_paragraph()


def render_appendix_usage(doc, usage):
    add_heading(doc, '附录E：使用说明', level=2)
    usage_table = doc.add_table(rows=len(usage) + 1, cols=2)
    for i, h in enumerate(['项目', '说明']):
        set_cell_text(usage_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, (k, v) in enumerate(usage):
        set_cell_text(usage_table.rows[ri + 1].cells[0], k, bold=True, size=10)
        set_cell_bg(usage_table.rows[ri + 1].cells[0], 'F2F2F2')
        set_cell_text(usage_table.rows[ri + 1].cells[1], v, size=10)
    set_table_widths(usage_table, [2.5, 13.5])
    style_header_row(usage_table)
    set_table_borders(usage_table)
    doc.add_paragraph()


def render_signature(doc, signers):
    add_heading(doc, '签字确认', level=1)
    sign_table = doc.add_table(rows=len(signers) + 1, cols=4)
    for i, h in enumerate(['角色', '姓名', '签字', '日期']):
        set_cell_text(sign_table.rows[0].cells[i], h, bold=True, size=11)
    for ri, (role, *_) in enumerate(signers):
        set_cell_text(sign_table.rows[ri + 1].cells[0], role, bold=True, size=11)
        for ci in range(1, 4):
            set_cell_text(sign_table.rows[ri + 1].cells[ci], '', size=11)
    set_table_widths(sign_table, [5.0, 3.5, 4.0, 3.5])
    style_header_row(sign_table)
    set_table_borders(sign_table)


# ============================================================
# 主函数
# ============================================================

def generate_survey_docx(config, output_path):
    """生成完整业务调研问卷 Word 文档。

    config 必填字段见 references/config_schema.md
    """
    doc = Document()

    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Microsoft YaHei'
    style_normal.font.size = Pt(11)
    style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    style_normal._element.rPr.rFonts.set(qn('w:ascii'), 'Microsoft YaHei')
    style_normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Microsoft YaHei')

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)

    render_cover(doc, config)
    render_attendees(doc, config['attending_dept'])
    render_instruction(doc, config['project_context'])
    render_toc(doc, config['toc_items'])
    render_overview(doc, config['overview_items'], config.get('core_upgrade'))
    render_current_status(doc, config['modules'])
    render_high_level_interview(doc, config['interview_topics'])
    render_it_systems(doc, config['it_systems'])
    render_data_interfaces(doc, config['data_flows'])
    render_reports(doc, config['reports'])
    render_appendix_resources(doc, config['resources'])
    render_appendix_sampling(doc, config.get('sampling_desc',
                              '建议抽查3~5个典型对象,覆盖主要业务领域。'))
    render_appendix_pain_points(doc)
    render_appendix_summary_dims(doc, config['summary_dims'])
    render_appendix_usage(doc, config['usage'])
    render_signature(doc, config['signers'])

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    print("此脚本为通用生成器。")
    print("请使用 generate_survey_docx(config, output_path) 调用。")
    print("详见 references/config_schema.md")