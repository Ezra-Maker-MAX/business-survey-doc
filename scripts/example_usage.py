"""
business-survey-docx skill 使用示例

演示如何调用 build_survey_docx.generate_survey_docx() 生成调研问卷。

运行方式（从本目录）:
    python example_usage.py
    或
    python example_usage.py --all
"""

import argparse
import os
import sys

# 添加 scripts / assets 目录到路径
_SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
_SKILL_DIR = os.path.dirname(_SCRIPT_DIR)
sys.path.insert(0, _SCRIPT_DIR)
sys.path.insert(0, os.path.join(_SKILL_DIR, 'assets'))

from build_survey_docx import generate_survey_docx


def example_minimal(output_path=None):
    """示例1：最小配置 - 手工指定少量问题（MES/生产管理场景演示）"""

    if output_path is None:
        output_path = os.path.join(_SKILL_DIR, 'output', 'example_mes_survey.docx')

    config = {
        'project_name': 'XX公司ERP实施项目',
        'doc_title': '生产管理业务调研问卷',
        'subtitle': '（V1.0｜2模块）',
        'info_rows': [
            ('项目名称', 'XX公司ERP实施项目', '', ''),
            ('文档编号', 'XX_V1.0_20260812', '', ''),
            ('版本', 'V1.0', '日期', '2026-08-12'),
            ('编制方', '咨询方', '面向客户', 'XX公司'),
            ('调研主题', '生产管理业务现状', '', ''),
            ('调研范围', '2 模块 4 个子项 / 9 条问题', '', ''),
        ],
        'attending_dept': '生产部、计划部、车间主任',
        'project_context': (
            'XX公司ERP实施项目旨在通过数字化转型提升生产管理效率。'
            '本次调研重点了解生产计划、车间执行、工序管理三大业务的现状、'
            '问题与需求,为后续ERP-MES蓝图设计提供输入。'
        ),

        'toc_items': [
            ('模块一', '生产计划管理', 'Q1-Q7', '7'),
            ('\u30001.1', '主计划管理', 'Q1-Q5', '5'),
            ('\u30001.2', '物料需求计划', 'Q6-Q7', '2'),
            ('模块二', '车间执行管理', 'Q8-Q9', '2'),
            ('\u30002.1', '工单管理', 'Q8-Q9', '2'),
        ],

        'overview_items': [
            '本次生产管理业务调研覆盖主计划、MRP、工单管理、工序报工四大子领域。',
            ([
                ('调研方法:', {'bold': True}),
                ('现场访谈 + 资料收集 + 系统演示,预计3个工作日完成。', {}),
            ], True),
        ],

        'modules': [
            {
                'sec_num': '5.1',
                'name': '生产计划管理',
                'intro': '生产计划是ERP/MES的核心,涵盖主生产计划(MPS)和物料需求计划(MRP)。',
                'subs': [
                    ('主计划管理', [
                        {'idx': 1, 'level': '必问', 'q': '主生产计划的编制频率和周期？计划周期长度？',
                         'follow': '计划变更的频次？', 'need': '主计划样例'},
                        {'idx': 2, 'level': '必问', 'q': '主计划的数据来源（订单/预测/库存）？',
                         'follow': '订单和预测的占比？', 'need': '数据源说明'},
                        {'idx': 3, 'level': '应问', 'q': '计划冻结区的设置规则？', 'need': '冻结规则文档'},
                        {'idx': 4, 'level': '深挖', 'q': '计划排程的约束条件（产能/物料/交期）如何平衡？',
                         'follow': '瓶颈工序如何识别？', 'need': '排程逻辑说明'},
                        {'idx': 5, 'level': '必问', 'q': '插单/改单的处理流程？对计划的冲击如何管理？',
                         'need': '变更管理流程'},
                    ]),
                    ('物料需求计划', [
                        {'idx': 6, 'level': '必问', 'q': 'MRP计算的频率和逻辑？', 'need': 'MRP逻辑说明'},
                        {'idx': 7, 'level': '应问', 'q': 'MRP异常的处理流程？', 'need': '异常处理流程'},
                    ]),
                ],
            },
            {
                'sec_num': '5.2',
                'name': '车间执行管理',
                'intro': '车间执行包括工单下达、报工、完工入库等业务。',
                'subs': [
                    ('工单管理', [
                        {'idx': 8, 'level': '必问', 'q': '工单的下达、变更、关闭流程？', 'need': '工单流程文档'},
                        {'idx': 9, 'level': '必问', 'q': '工单与生产订单的对应关系？', 'need': '工单样例'},
                    ]),
                ],
            },
        ],

        'interview_topics': [
            ('1', '战略定位', '生产管理在公司的战略定位...'),
            ('2', '数字化期望', '对MES系统的期望...'),
        ],

        'it_systems': [
            ('ERP系统', '', '', '', ''),
            ('MES系统', '', '', '', ''),
            ('WMS系统', '', '', '', ''),
        ],

        'data_flows': [
            ('ERP ↔ MES', '', '', ''),
            ('MES ↔ WMS', '', '', ''),
        ],

        'reports': [
            ('生产计划报表', '', '', ''),
            ('车间执行报表', '', '', ''),
        ],

        'resources': [
            (1, '生产计划', '主计划样例', '高'),
            (2, '工艺路线', '工艺路线文档', '高'),
        ],

        'summary_dims': [
            '1. 生产计划管理',
            '2. 车间执行管理',
            '3. 系统集成',
            '4. 核心痛点 TOP10',
        ],

        'usage': [
            ('调研项目', '生产管理业务调研'),
            ('调研范围', '2 模块 6 个子项'),
            ('调研原则', '先问现状,再看现场'),
        ],

        'signers': [
            ('甲方 业务负责人', '', '', ''),
            ('乙方 项目经理', '', '', ''),
        ],
    }

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    generate_survey_docx(config, output_path)
    print(f'✅ 示例1（MES最小配置）生成成功: {output_path}')
    return output_path


def example_eam_v4(output_path=None):
    """示例2：使用 V4.0 EAM 完整问题库生成 280 条调研问卷"""

    from eam_v4_data import build_eam_v4_config

    if output_path is None:
        output_path = os.path.join(_SKILL_DIR, 'output', 'example_eam_v4_survey.docx')

    config, questions = build_eam_v4_config()

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    generate_survey_docx(config, output_path)
    print(f'✅ 示例2（EAM V4.0完整版）生成成功: {output_path}')
    print(f'   问题总数: {len(questions)}')
    return output_path


def verify_output(output_path):
    """生成后简单校验"""
    from docx import Document
    import re

    doc = Document(output_path)
    q_indices = []
    for p in doc.paragraphs:
        text = p.text.strip()
        m = re.match(r'Q(\d+)', text)
        if m:
            q_indices.append(int(m.group(1)))

    q_indices.sort()
    if not q_indices:
        print('⚠️ 未找到任何问题！')
        return False

    expected = list(range(1, max(q_indices) + 1))
    missing = set(expected) - set(q_indices)
    print(f'   文档段落数: {len(doc.paragraphs)}')
    print(f'   表格数: {len(doc.tables)}')
    print(f'   问题数: {len(q_indices)} (Q{min(q_indices)}-Q{max(q_indices)})')
    if missing:
        print(f'   ⚠️ 缺失编号: {sorted(missing)}')
        return False
    print('   ✅ Q编号连续无缺失')
    return True


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='business-survey-docx 示例')
    parser.add_argument('--all', action='store_true', help='运行所有示例')
    parser.add_argument('--eam', action='store_true', help='仅运行EAM V4.0完整示例')
    parser.add_argument('--output', type=str, default=None, help='输出文件路径（EAM示例）')
    args = parser.parse_args()

    if args.eam:
        path = example_eam_v4(args.output)
        ok = verify_output(path)
        print(f'验证结果: {"通过" if ok else "失败"}')
    elif args.all:
        p1 = example_minimal()
        p2 = example_eam_v4()
        print('\n=== 验证 ===')
        verify_output(p1)
        verify_output(p2)
    else:
        print('用法: python example_usage.py [--all | --eam] [--output PATH]')
        print('  --all    运行所有示例')
        print('  --eam    仅运行EAM V4.0完整版示例')
        print()
        print('快速体验: python example_usage.py --eam')