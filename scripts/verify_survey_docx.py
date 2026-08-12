"""
调研问卷 docx 验证脚本

自动检查生成的调研问卷文档是否满足标准格式：
1. 章节完整性（9章正文 + 5附录 + 签字栏）
2. Q编号连续性（从1开始连续，无缺失无重复）
3. 问题总数
4. 表格数量
5. 问题格式抽查（Q编号+[等级]）

用法:
    python verify_survey_docx.py <path/to/survey.docx>
    或
    python verify_survey_docx.py --strict <path/to/survey.docx>  # 严格模式
"""

import argparse
import re
import sys

from docx import Document


# 标准章节清单
REQUIRED_SECTIONS = [
    '一、参会人员',
    '二、调研说明',
    '三、调研目录',
    '四、总体概述',
    '五、现状描述',
    '六、高层访谈',
    '七、IT系统',
    '八、数据接口',
    '九、标准报表',
    '附录A：资料索取清单',
    '附录B：典型对象现场抽查',
    '附录C：痛点需求池',
    '附录D：调研总结',
    '附录E：使用说明',
    '签字确认',
]

VALID_LEVELS = {'必问', '应问', '深挖'}


def verify_survey_docx(docx_path, strict=False):
    """
    验证调研问卷文档。

    返回 (ok: bool, report: list[str])
    """
    report = []
    ok = True

    try:
        doc = Document(docx_path)
    except Exception as e:
        return False, [f'❌ 无法打开文档: {e}']

    # 1. 提取所有段落实文本
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    # 2. 章节完整性检查
    found_sections = []
    for text in paragraphs:
        for section in REQUIRED_SECTIONS:
            if text == section or text.startswith(section):
                if section not in found_sections:
                    found_sections.append(section)

    missing_sections = [s for s in REQUIRED_SECTIONS if s not in found_sections]
    if missing_sections:
        ok = False
        report.append(f'❌ 缺失章节: {", ".join(missing_sections)}')
    else:
        report.append(f'✅ 章节完整（{len(REQUIRED_SECTIONS)}/15）')

    # 3. Q编号提取与连续性检查
    q_indices = []
    q_level_issues = []
    for text in paragraphs:
        m = re.match(r'Q(\d+)\s+\[(.+?)\]', text)
        if m:
            idx = int(m.group(1))
            level = m.group(2)
            q_indices.append(idx)
            if level not in VALID_LEVELS:
                q_level_issues.append((idx, level))

    q_indices.sort()

    if not q_indices:
        ok = False
        report.append('❌ 未找到任何调研问题（Q编号）')
    else:
        # 连续性
        expected = list(range(1, max(q_indices) + 1))
        missing_ids = [i for i in expected if i not in set(q_indices)]
        duplicates = [i for i in set(q_indices) if q_indices.count(i) > 1]

        report.append(f'ℹ️ 问题总数: {len(q_indices)} (Q{min(q_indices)}-Q{max(q_indices)})')

        if missing_ids:
            ok = False
            report.append(f'❌ 缺失Q编号: {missing_ids[:20]}{"..." if len(missing_ids) > 20 else ""}')
        else:
            report.append('✅ Q编号连续无缺失')

        if duplicates:
            ok = False
            report.append(f'❌ 重复Q编号: {duplicates}')

        # 问题等级合法性
        if q_level_issues:
            ok = False
            report.append(f'❌ 非法等级: {q_level_issues[:10]}')

    # 4. 表格数量
    table_count = len(doc.tables)
    report.append(f'ℹ️ 表格数: {table_count}')
    if strict and table_count < 8:
        ok = False
        report.append(f'❌ 严格模式：表格数过少（<8）')

    # 5. 完整性问题抽查（每条问题应有"答复"行）
    if strict:
        answer_count = sum(1 for t in paragraphs if t == '答复:' or t.startswith('答复:'))
        report.append(f'ℹ️ 答复行数: {answer_count} (问题数: {len(q_indices)})')
        if answer_count < len(q_indices) * 0.9:
            ok = False
            report.append(f'❌ 严格模式：答复行数不足（{answer_count}/{len(q_indices)}）')

    # 6. 统计汇总
    if ok:
        report.append(f'🎉 验证通过！文档结构符合 business-survey-docx 标准。')
    else:
        report.append(f'⚠️ 验证未通过，请检查上述问题。')

    return ok, report


def main():
    parser = argparse.ArgumentParser(description='验证调研问卷docx')
    parser.add_argument('docx_path', help='调研问卷docx路径')
    parser.add_argument('--strict', action='store_true', help='严格模式（额外检查答复行数等）')
    args = parser.parse_args()

    ok, report = verify_survey_docx(args.docx_path, strict=args.strict)
    print(f'\n=== 验证报告: {args.docx_path} ===')
    for line in report:
        print(f'  {line}')
    print()
    sys.exit(0 if ok else 1)


if __name__ == '__main__':
    main()